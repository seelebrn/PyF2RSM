import os
import csv
from docx import Document as DocxDocument
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from config import CHUNK_SIZE, CHUNK_OVERLAP, config

def extract_text_from_docx(file_path):
    """Extracts and returns text from a .docx file."""
    doc = DocxDocument(file_path)
    text = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():  # Only include non-empty paragraphs
            text.append(paragraph.text.strip())
    return "\n".join(text)

def extract_text_from_csv(file_path):
    """Extracts and returns text from a .csv file, combining post title, date, and content."""
    text = []
    with open(file_path, "r", encoding="utf-8") as f:
        datareader = csv.reader(f, delimiter=";")
        next(datareader)  # Skip the header row
        for row in datareader:
            if len(row) > 3:
                cam_date = row[0].strip() #post title
                cam_ref = row[1].strip() #post date & hour
                cam_asker = row[2].strip() #post content
                cam_subject = row[3].strip() #post content

                post_title = row[1].strip() #post title
                post_date = row[9].strip() #post date & hour
                post_content = row[3].strip() #post content
                #Combine the 3 columns into one structured string
                combined_text = f"Title: {post_title}\nDate: {post_date}\nContent: {post_content}"


                text.append(combined_text)
    return "\n\n".join(text)  # Join posts with a double newline for separation


def load_and_split_documents():
    """Loads and splits text from all .csv and .docx files in a folder."""
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)
    documents = []

    # Process all files in the base folder
    for file_name in os.listdir(config["BASE_FOLDER_PATH"]):
        file_path = os.path.join(config["BASE_FOLDER_PATH"], file_name)
        if file_name.endswith(".docx"):
            print(f"Processing DOCX file: {file_path}")
            full_text = extract_text_from_docx(file_path)
        elif file_name.endswith(".csv"):
            print(f"Processing CSV file: {file_path}")
            full_text = extract_text_from_csv(file_path)
        else:
            print(f"Skipping unsupported file: {file_name}")
            continue

        # Split text into chunks
        chunks = text_splitter.split_text(full_text)
        documents.extend(chunks)

    print(f"Total chunks created: {len(documents)}")
    return [Document(page_content=chunk) for chunk in documents]
